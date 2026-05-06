from __future__ import annotations

import io
from typing import Literal

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.services.datasets import StoredDataset, profile_dataset


def export_dataset(dataset: StoredDataset, format_name: str) -> tuple[bytes, str, str]:
    frame = dataset.frame.copy()
    if format_name in {"csv", "r", "spss", "gis", "stata"}:
        buffer = io.StringIO()
        frame.to_csv(buffer, index=False)
        extension = "csv"
        media = "text/csv"
        if format_name == "r":
            frame.attrs["note"] = "R-ready CSV with UTF-8 headers and analysis-safe column names."
        return buffer.getvalue().encode("utf-8"), media, extension
    if format_name == "tsv":
        buffer = io.StringIO()
        frame.to_csv(buffer, index=False, sep="\t")
        return buffer.getvalue().encode("utf-8"), "text/tab-separated-values", "tsv"
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        frame.to_excel(writer, sheet_name="clean_data", index=False)
        pd.DataFrame({"cleaning_log": dataset.log}).to_excel(writer, sheet_name="cleaning_log", index=False)
        dictionary = pd.DataFrame(
            {
                "variable": frame.columns,
                "label": [column.replace("_", " ").title() for column in frame.columns],
                "role": ["GIS join key" if "district" in column.lower() or "code" in column.lower() else "analysis variable" for column in frame.columns],
            }
        )
        dictionary.to_excel(writer, sheet_name="variable_dictionary", index=False)
    return output.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"


def metadata_document(dataset_id: str, dataset: StoredDataset, kind: Literal["pdf", "docx"]) -> tuple[bytes, str, str]:
    profile = profile_dataset(dataset_id, dataset)
    if kind == "docx":
        doc = Document()
        doc.add_heading("GeoClean Research Studio Metadata Report", level=1)
        doc.add_paragraph(f"Dataset: {dataset.file_name}")
        doc.add_paragraph(f"Rows: {profile.file.rows}; Columns: {profile.file.columns}; Quality score: {profile.quality_score}%")
        doc.add_heading("Cleaning Log", level=2)
        for item in dataset.log:
            doc.add_paragraph(item, style="List Bullet")
        doc.add_heading("Variable Dictionary", level=2)
        table = doc.add_table(rows=1, cols=4)
        headers = table.rows[0].cells
        headers[0].text = "Variable"
        headers[1].text = "Type"
        headers[2].text = "Missing"
        headers[3].text = "Unique"
        for column in profile.columns:
            cells = table.add_row().cells
            cells[0].text = column.name
            cells[1].text = column.inferred_type
            cells[2].text = str(column.missing_count)
            cells[3].text = str(column.unique_count)
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    y = 750
    pdf.setFont("Helvetica-Bold", 15)
    pdf.drawString(48, y, "GeoClean Research Studio Metadata Report")
    y -= 32
    pdf.setFont("Helvetica", 10)
    lines = [
        f"Dataset: {dataset.file_name}",
        f"Rows: {profile.file.rows}; Columns: {profile.file.columns}; Quality score: {profile.quality_score}%",
        "Cleaning Log:",
        *[f"- {item}" for item in dataset.log[:12]],
        "Variable Dictionary:",
        *[f"- {column.name}: {column.inferred_type}, missing={column.missing_count}, unique={column.unique_count}" for column in profile.columns[:24]],
    ]
    for line in lines:
        if y < 60:
            pdf.showPage()
            y = 750
            pdf.setFont("Helvetica", 10)
        pdf.drawString(48, y, line[:110])
        y -= 18
    pdf.save()
    return output.getvalue(), "application/pdf", "pdf"
